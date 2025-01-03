import docx
import os
import pandas as pd
import random
from openpyxl import load_workbook
import re
from collections import Counter

test_id = 3
test_offset = 6 + test_id
report_dir = f'./22B4-实训报告3(附件)'
excel_dir = './B4班过程性考核成绩登记表.xlsx'

name_list = []


def extract_nested_table(table):
    for row in table.rows:
        for cell in row.cells:
            if cell.tables:  # 检查单元格中是否包含嵌套表格
                return cell.tables
                
    return None
    


def extract_nested_table_text(table):
    nested_table_data = []
    for nested_table in table:
        for nested_row in nested_table.rows:
            for nested_cell in nested_row.cells:
                nested_table_data.append(nested_cell.text)
    return nested_table_data


def input_score(table, start, step, sheet):
    count = 0
    pos = start
    for nested_table in table:
        for nested_row in nested_table.rows:
            for i, nested_cell in enumerate(nested_row.cells):  
                if count == pos:
                    if nested_cell.text.isdigit() or nested_cell.text == '':
                        # 生成75到81之间的随机数（包含75和81） 
                        random_number = random.randint(75, 81) - 0
                        if nested_row.cells[i-1].text != '':
                            self_score = int(nested_row.cells[i-1].text)
                            total_score = self_score*0.2 + random_number*0.8
                            # print(total_score)
                            nested_cell.text = str(random_number)
                            nested_row.cells[i+1].text = str(round(total_score))

                            # 使用正则表达式匹配中文字符 
                            chinese_text = re.findall(r'[\u4e00-\u9fff]+', nested_row.cells[i-3].text) 
                            # 将匹配到的字符连接成字符串 
                            target_name = ''.join(chinese_text)
                            if target_name in name_list:
                                return -1
                            else:
                                name_list.append(target_name)
                            print(target_name)

                            # 查找姓名并返回位置
                            found = False
                            for row in sheet.iter_rows():
                                for cell in row:
                                    if cell.value == target_name:
                                        print(f"找到 {target_name} 在单元格: {cell.coordinate}")
                                        found = True

                                        target_cell = cell
                                        # 获取目标单元格的行和列 
                                        row_idx = target_cell.row 
                                        col_idx = target_cell.column
                                        # print(row_idx, col_idx)

                                        test1_cell = sheet.cell(row=row_idx, column=col_idx + test_offset)
                                        test1_cell.value = str(round(total_score))

                                        break
                                if found:
                                    break
                            if not found:
                                print(f"未找到 {target_name}")

                        pos += step
                count += 1
 
    return 0
                                

def main():

    # 加载现有的Excel文件
    workbook = load_workbook(excel_dir)
    # 选择工作表
    sheet = workbook['Sheet1']  # 替换为你需要的工作表名称

    # 遍历当前目录下的所有文件 
    for filename in os.listdir(report_dir): 
        if filename.endswith('.docx'): 
            file_path = os.path.join(report_dir, filename) 
            print(file_path)
            doc = docx.Document(file_path)
            # 遍历文档中的所有表格
            for table in doc.tables:
                n_table = extract_nested_table(table)
                nested_text = extract_nested_table_text(n_table)
                if nested_text:  # 如果有嵌套表格内容，打印出来
                    print("嵌套表格内容:")
                    print(nested_text)
            


            # 遍历文档中的所有表格
            for table in doc.tables:
                n_table = extract_nested_table(table)
                ret = input_score(n_table, 8, 5, sheet)
            
            if ret == 0:
                # 保存文档 
                doc.save(file_path)
            else:
                # 删除重复文档
                os.remove(file_path)
                continue

            doc = docx.Document(file_path)
            # 遍历文档中的所有表格
            for table in doc.tables:
                n_table = extract_nested_table(table)
                nested_text = extract_nested_table_text(n_table)
                if nested_text:  # 如果有嵌套表格内容，打印出来
                    print("嵌套表格内容:")
                    print(nested_text)

    workbook.save(excel_dir)

    element_count = Counter(name_list)
    print(element_count)


if __name__ == "__main__":
    main()

